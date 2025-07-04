from fastapi import FastAPI, UploadFile, File, Query, Path, Body, Form
from fastapi.responses import JSONResponse, StreamingResponse
from fastapi import APIRouter
from fastapi.middleware.cors import CORSMiddleware
from typing import List, Optional, Dict, Any
import psycopg2
import os
import joblib
import numpy as np
import docx
import re
import io
import zipfile
from app.extractors import extract_text_from_pdf  # Adjust import as needed

app = FastAPI()
router = APIRouter()

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Or specify your UI's URL for more security
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

DB_CONFIG = {
    "host": "10.30.0.102",
    "port": 5432,
    "database": "vectordb",
    "user": "devuser",
    "password": "StrongPassword123"
}

SUPPORTED_TYPES = {".pdf", ".doc", ".docx", ".txt"}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

def get_db_conn():
    return psycopg2.connect(**DB_CONFIG)

def parse_resume(file_bytes, filename):
    name = os.path.splitext(filename)[0].replace("_", " ").title()
    return {
        "name": name,
        "email": f"{name.lower().replace(' ', '')}@example.com",
        "skills": ["Python", "React"]
    }

def extract_text(file_bytes, filename):
    ext = os.path.splitext(filename)[1].lower()
    temp_path = f"temp_upload{ext}"
    with open(temp_path, "wb") as f:
        f.write(file_bytes)
    try:
        if ext == ".pdf":
            text = extract_text_from_pdf(temp_path)
        elif ext == ".docx":
            doc = docx.Document(temp_path)
            text = "\n".join([para.text for para in doc.paragraphs])
        elif ext == ".txt":
            with open(temp_path, "r", encoding="utf-8", errors="ignore") as f2:
                text = f2.read()
        else:
            text = ""
    finally:
        os.remove(temp_path)
    return text

def extract_experience_years(text: str):
    # Look for patterns like "5 years", "3+ years", "2-years", etc.
    match = re.search(r'(\d+)\s*\+?\s*[-]?\s*years?', text, re.IGNORECASE)
    if match:
        return int(match.group(1))
    return "Fresher"

@app.post("/api/resumes/upload")
async def upload_resumes(files: List[UploadFile] = File(...)):
    print("Upload endpoint called")
    results = []
    failed_details = []

    # Load embedding models once
    tfidf = joblib.load("embeddings/tfidf.joblib")
    svd = joblib.load("embeddings/svd.joblib")

    conn = get_db_conn()
    cursor = conn.cursor()

    for file in files:
        file_result = {"filename": file.filename}
        ext = os.path.splitext(file.filename)[1].lower()

        # Validation: file type
        if ext not in SUPPORTED_TYPES:
            file_result["upload_status"] = "failed"
            file_result["error"] = "Unsupported file format"
            failed_details.append({"file": file.filename, "reason": "Unsupported file type"})
            results.append(file_result)
            continue

        contents = await file.read()
        # Validation: file size
        if len(contents) > MAX_FILE_SIZE:
            file_result["upload_status"] = "failed"
            file_result["error"] = "File size exceeds 10MB"
            failed_details.append({"file": file.filename, "reason": "File size exceeds 10MB"})
            results.append(file_result)
            continue

        try:
            # Save file to pdf_file table
            cursor.execute(
                "INSERT INTO pdf_file (filename, content) VALUES (%s, %s) RETURNING id",
                (file.filename, psycopg2.Binary(contents))
            )
            resume_id = cursor.fetchone()[0]
            parsed_data = parse_resume(contents, file.filename)

            # Extract text for embedding
            text = extract_text(contents, file.filename)
            if text.strip():
                tfidf_vec = tfidf.transform([text])
                svd_vec = svd.transform(tfidf_vec)
                embedding = svd_vec[0].tolist()  # VECTOR(300) expects a list/array

                # Save to resume_embeddings table with pdf_id
                cursor.execute(
                    """
                    INSERT INTO resume_embeddings (pdf_id, name, resume_text, embedding, file_data, file_name)
                    VALUES (%s, %s, %s, %s, %s, %s)
                    """,
                    (
                        resume_id,
                        parsed_data["name"],
                        text,
                        embedding,
                        psycopg2.Binary(contents),
                        file.filename
                    )
                )

            file_result.update({
                "id": f"resume_{resume_id}",
                "upload_status": "success",
                "parsed_data": parsed_data
            })
        except Exception as e:
            print(f"Error processing file {file.filename}: {e}")
            conn.rollback()
            file_result["upload_status"] = "failed"
            file_result["error"] = str(e)
            failed_details.append({"file": file.filename, "reason": str(e)})
        results.append(file_result)

    conn.commit()
    cursor.close()
    conn.close()

    uploaded_count = sum(1 for r in results if r.get("upload_status") == "success")
    failed_count = sum(1 for r in results if r.get("upload_status") == "failed")

    # If all failed, return error response
    if uploaded_count == 0:
        return JSONResponse({
            "success": False,
            "message": "Resume upload failed due to validation error",
            "error": {
                "code": "VALIDATION_ERROR",
                "details": failed_details
            }
        })

    # Otherwise, return normal response
    return JSONResponse({
        "success": True,
        "message": "Files uploaded successfully",
        "data": {
            "uploaded_count": uploaded_count,
            "failed_count": failed_count,
            "resumes": results
        }
    })

@app.get("/api/resumes")
def get_resumes(
    page: int = Query(1, ge=1),
    limit: int = Query(10, ge=1, le=100)
):
    conn = get_db_conn()
    cursor = conn.cursor()

    offset = (page - 1) * limit

    # Get total count
    cursor.execute("SELECT COUNT(*) FROM pdf_file")
    total_items = cursor.fetchone()[0]
    total_pages = (total_items + limit - 1) // limit

    # Fetch resumes with pagination
    cursor.execute(
        """
        SELECT id, filename, created_at
        FROM pdf_file
        ORDER BY created_at DESC
        LIMIT %s OFFSET %s
        """,
        (limit, offset)
    )
    rows = cursor.fetchall()

    resumes = []
    for row in rows:
        resume_id, filename, upload_date = row

        # Fetch parsed data and resume_text from resume_embeddings if available
        cursor.execute(
            """
            SELECT name, resume_text
            FROM resume_embeddings
            WHERE file_name = %s
            LIMIT 1
            """,
            (filename,)
        )
        parsed = cursor.fetchone()
        if parsed:
            name, resume_text = parsed
            experience_years = extract_experience_years(resume_text or "")
            # Extract summary/career objective up to (but not including) Education section
            def extract_resume_summary(text):
                import re
                if not text:
                    return None
                summary_patterns = [
                    r'(?:summary|career objective|professional summary|profile|objective statement)\s*[:\-\n]+(.{0,1000})',
                ]
                for pat in summary_patterns:
                    m = re.search(pat, text, re.IGNORECASE | re.DOTALL)
                    if m:
                        summary = m.group(1)
                        summary = re.split(r'\n+\s*(Education|EDUCATION|Academic|ACADEMIC)\b', summary)[0]
                        return summary.strip()
                summary = re.split(r'\n+\s*(Education|EDUCATION|Academic|ACADEMIC)\b', text)[0]
                return summary.strip()[:400]

            # Extract skills using pyresparser if available, else fallback to regex
            def extract_skills(text):
                try:
                    from pyresparser import ResumeParser
                    import tempfile
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.txt', mode='w', encoding='utf-8') as temp:
                        temp.write(text)
                        temp_path = temp.name
                    data = ResumeParser(temp_path).get_extracted_data()
                    os.remove(temp_path)
                    skills = data.get('skills')
                    if skills:
                        return skills
                except Exception as e:
                    print(f"Skill extraction failed: {e}")
                # Fallback: simple regex for common skills
                common_skills = [
                    'python', 'java', 'c++', 'c#', 'javascript', 'react', 'node', 'sql', 'html', 'css', 'aws', 'docker',
                    'django', 'flask', 'angular', 'typescript', 'git', 'linux', 'excel', 'pandas', 'numpy', 'machine learning',
                    'deep learning', 'nlp', 'tensorflow', 'keras', 'pytorch', 'tableau', 'power bi', 'spark', 'hadoop', 'scala',
                    'php', 'ruby', 'go', 'swift', 'kotlin', 'android', 'ios', 'rest', 'graphql', 'mongodb', 'postgresql', 'mysql'
                ]
                found = set()
                for skill in common_skills:
                    if re.search(r'\b' + re.escape(skill) + r'\b', text, re.IGNORECASE):
                        found.add(skill.title())
                return sorted(found)

            description = extract_resume_summary(resume_text)
            skills = extract_skills(resume_text or "")
            # Extract email from resume text
            def extract_email(text):
                import re
                if not text:
                    return None
                matches = re.findall(r'[\w\.-]+@[\w\.-]+\.\w+', text)
                return matches[0] if matches else None

            email = extract_email(resume_text or "")
            parsed_data = {
                "name": name,
                "email": email,
                "skills": skills,
                "experience_years": experience_years,
                "description": description
            }
        else:
            parsed_data = {
                "name": None,
                "email": None,
                "skills": [],
                "experience_years": "Fresher"
            }

        resumes.append({
            "id": f"resume_{resume_id}",
            "filename": filename,
            "upload_date": upload_date.isoformat() if upload_date else None,
            "parsed_data": parsed_data,
            "match_score": None
        })

    cursor.close()
    conn.close()

    return {
        "success": True,
        "data": {
            "resumes": resumes,
            "pagination": {
                "current_page": page,
                "total_pages": total_pages,
                "total_items": total_items,
                "items_per_page": limit
            }
        }
    }

@app.get("/api/resumes/{resume_id}/download")
def download_resume(
    resume_id: str = Path(..., description="Resume ID, e.g., resume_1"),
    format: str = Query("pdf", description="Download format (default: pdf)")
):
    # Extract numeric ID from "resume_1"
    try:
        db_id = int(resume_id.replace("resume_", ""))
    except Exception:
        return JSONResponse(
            status_code=400,
            content={"success": False, "message": "Invalid resume ID format"}
        )

    conn = get_db_conn()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT filename, content FROM pdf_file WHERE id = %s",
        (db_id,)
    )
    row = cursor.fetchone()
    cursor.close()
    conn.close()

    if not row:
        return JSONResponse(
            status_code=404,
            content={"success": False, "message": "Resume not found"}
        )

    filename, file_data = row
    # You can add logic for other formats if needed
    content_type = "application/pdf" if filename.lower().endswith(".pdf") else "application/octet-stream"
    return StreamingResponse(
        iter([file_data]),
        media_type=content_type,
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"'
        }
    )

@app.post("/api/resumes/download-all")
async def download_all_resumes(
    body: dict = Body(...),
    format: str = Query("zip", description="Download format (default: zip)")
):
    resume_ids = body.get("resume_ids", None)

    conn = get_db_conn()
    cursor = conn.cursor()

    # If resume_ids is provided, filter by those IDs; else, get all
    if resume_ids:
        db_ids = []
        for rid in resume_ids:
            try:
                db_ids.append(int(rid.replace("resume_", "")))
            except Exception:
                continue
        if not db_ids:
            cursor.close()
            conn.close()
            return JSONResponse(
                status_code=400,
                content={"success": False, "message": "No valid resume IDs provided"}
            )
        sql = "SELECT filename, content FROM pdf_file WHERE id = ANY(%s)"
        cursor.execute(sql, (db_ids,))
    else:
        cursor.execute("SELECT filename, content FROM pdf_file")
    rows = cursor.fetchall()
    cursor.close()
    conn.close()

    if not rows:
        return JSONResponse(
            status_code=404,
            content={"success": False, "message": "No resumes found"}
        )

    # Create ZIP in memory
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
        for filename, file_data in rows:
            zipf.writestr(filename, file_data)
    zip_buffer.seek(0)

    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={
            "Content-Disposition": 'attachment; filename="all_resumes.zip"'
        }
    )

def extract_jd_requirements(jd_text: str) -> dict:
    # Dummy extraction logic; replace with real NLP as needed
    return {
        "required_skills": ["Python", "Django", "SQL"],
        "preferred_skills": ["Docker", "AWS"],
        "experience_years": 5,
        "education": "Bachelor's degree in Computer Science"
    }

def get_job_category_and_seniority(jd_text: str) -> (str, str):
    # Dummy logic; replace with real classification if needed
    return "Software Development", "Senior"

@router.post("/api/jobs/process-text-and-match")
async def process_text_and_match(
    body: Dict[str, Any] = Body(...)
):
    jd_text = body.get("job_description")
    title = body.get("title", "Job")
    resume_ids = body.get("resume_ids", None)

    # 1. Extract requirements and meta info from JD
    requirements = extract_jd_requirements(jd_text)
    job_category, seniority_level = get_job_category_and_seniority(jd_text)

    # 2. Generate JD embedding
    tfidf = joblib.load("embeddings/tfidf.joblib")
    svd = joblib.load("embeddings/svd.joblib")
    jd_vec = svd.transform(tfidf.transform([jd_text]))[0]

    # 3. Fetch resumes to match
    conn = get_db_conn()
    cursor = conn.cursor()
    if resume_ids:
        db_ids = [int(rid.replace("resume_", "")) for rid in resume_ids]
        cursor.execute("SELECT id, file_name, resume_text, embedding, name FROM resume_embeddings WHERE id = ANY(%s)", (db_ids,))
    else:
        cursor.execute("SELECT id, file_name, resume_text, embedding, name FROM resume_embeddings")
    resumes = cursor.fetchall()
    cursor.close()
    conn.close()

    matched_resumes = []
    total_score = 0
    for rid, filename, resume_text, embedding, name in resumes:
        # Convert embedding to numpy array
        if isinstance(embedding, str):
            emb = np.array(eval(embedding))
        else:
            emb = np.array(embedding, dtype=np.float32)
        if emb.shape != jd_vec.shape:
            continue  # skip if shape mismatch
        sim = float(np.dot(jd_vec, emb) / (np.linalg.norm(jd_vec) * np.linalg.norm(emb)))
        match_score = round(sim * 100, 1)
        total_score += match_score

        # Dummy skill matching logic
        resume_skills = ["Python", "React", "SQL", "JavaScript"]

        required_skills = requirements["required_skills"]
        matching_skills = [s for s in required_skills if s in resume_skills]
        missing_skills = [s for s in required_skills if s not in resume_skills]

        # Skill match percentage
        if required_skills:
            skills_match_pct = int(len(matching_skills) / len(required_skills) * 100)
        else:
            skills_match_pct = 0

        # Overall fit logic
        if skills_match_pct >= 80:
            overall_fit = "Excellent"
        elif skills_match_pct >= 70:
            overall_fit = "Good"
        else:
            overall_fit = "Average"

        match_details = {
            "skills_match": skills_match_pct,
            "experience_match": 80,  # You can add real logic here
            "overall_fit": overall_fit
        }

        parsed_data = {
            "name": name,
            "email": f"{name.lower().replace(' ', '')}@example.com",
            "skills": resume_skills,
            "experience_years": requirements["experience_years"]
        }

        matched_resumes.append({
            "id": f"resume_{rid}",
            "filename": filename,
            "match_score": match_score,
            "match_details": match_details,
            "parsed_data": parsed_data,
            "missing_skills": missing_skills,
            "matching_skills": matching_skills
        })

    avg_score = round(total_score / len(matched_resumes), 1) if matched_resumes else 0

    return {
        "success": True,
        "message": "Job description processed and resumes matched successfully",
        "data": {
            "job_analysis": {
                "title": title,
                "processed_description": jd_text,
                "extracted_requirements": requirements,
                "job_category": job_category,
                "seniority_level": seniority_level
            },
            "matched_resumes": matched_resumes,
            "total_matched": len(matched_resumes),
            "average_score": avg_score
        }
    }

@router.post("/api/jobs/process-file-and-match")
async def process_file_and_match(
    file: UploadFile = File(...),
    title: Optional[str] = Form(None),
    resume_ids: Optional[str] = Form(None)
):
    # Save file temporarily
    ext = os.path.splitext(file.filename)[1].lower()
    temp_path = f"temp_jd{ext}"
    with open(temp_path, "wb") as f:
        f.write(await file.read())

    # Extract text
    if ext == ".pdf":
        jd_text = extract_text_from_pdf(temp_path)
    elif ext == ".docx":
        doc = docx.Document(temp_path)
        jd_text = "\n".join([para.text for para in doc.paragraphs])
    elif ext == ".txt":
        with open(temp_path, "r", encoding="utf-8", errors="ignore") as f2:
            jd_text = f2.read()
    else:
        jd_text = ""
    os.remove(temp_path)

    # Parse resume_ids if provided
    resume_ids_list = None
    if resume_ids:
        import json
        resume_ids_list = json.loads(resume_ids)

    # Call the text-based endpoint logic
    return await process_text_and_match({
        "job_description": jd_text,
        "title": title or "Job",
        "resume_ids": resume_ids_list
    })

app.include_router(router)